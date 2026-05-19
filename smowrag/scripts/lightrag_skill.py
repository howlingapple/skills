#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LightRAG Skill Implementation (异步版本 + 资源保护)

使用 LightRAG 框架实现知识图谱 RAG，替代自研 GraphRAG。
官方强烈建议使用 async 版本的接口 (ainsert, aquery)。

资源保护特性：
- LLM 并发信号量控制
- 指数退避重试机制
- Embedding 批处理优化
- API 超时保护
"""

import asyncio
import os
import sys
import time
import importlib
from pathlib import Path
from typing import Optional, List, Dict, Any, Callable
from functools import wraps

import numpy as np

# 优先使用本地 lightrag 源码（必须在 import lightrag 之前设置）
SKILL_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(SKILL_DIR))

from lightrag import LightRAG, QueryParam
from lightrag.utils import EmbeddingFunc

# 尝试导入 tenacity，如未安装则提供简单实现
try:
    from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
    HAS_TENACITY = True
except ImportError:
    HAS_TENACITY = False
    print("[警告] 未安装 tenacity，使用简单重试机制。建议: pip install tenacity")


def parse_pdf_simple(file_path: str) -> str:
    """PDF 解析：使用 PyPDF2（可作为配置中的 parser 引用）"""
    import PyPDF2
    text = ""
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

# =================================================================
# 资源保护：信号量控制并发
# =================================================================

# LLM 并发控制
LLM_SEMAPHORE = asyncio.Semaphore(2)

# Embedding 并发控制
EMBED_SEMAPHORE = asyncio.Semaphore(4)

# 网络请求超时配置
LLM_TIMEOUT = 600  # LLM 请求超时（秒），实体提取 prompt 很长需要更久
EMBED_TIMEOUT = 120  # Embedding 请求超时（秒）


def _load_skill_config() -> dict:
    """加载 skill 专属配置"""
    import yaml
    config_path = SKILL_DIR / "config.yaml"
    with open(config_path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def async_retry(max_attempts=5, min_wait=2, max_wait=60):
    """异步重试装饰器（简单实现，兼容无 tenacity 的情况）"""
    def decorator(func):
        if HAS_TENACITY:
            # 使用 tenacity 的高级重试
            return retry(
                stop=stop_after_attempt(max_attempts),
                wait=wait_exponential(multiplier=1, min=min_wait, max=max_wait),
                retry=retry_if_exception_type((Exception,)),
                reraise=True
            )(func)
        else:
            # 简单重试实现
            @wraps(func)
            async def wrapper(*args, **kwargs):
                last_exception = None
                for attempt in range(max_attempts):
                    try:
                        return await func(*args, **kwargs)
                    except Exception as e:
                        last_exception = e
                        wait_time = min(2 ** attempt, max_wait)
                        print(f"[重试] {func.__name__} 第 {attempt + 1}/{max_attempts} 次失败，{wait_time}s 后重试: {e}")
                        await asyncio.sleep(wait_time)
                raise last_exception
            return wrapper
    return decorator


# =================================================================
# 模块级文件读取函数（CLI 和 LightRAGSkill 共用）
# =================================================================

async def parse_pdf_with_parser(pdf_path: Path, parser: Callable) -> str:
    """使用给定的解析器解析 PDF"""
    result = parser(str(pdf_path))
    if asyncio.iscoroutine(result):
        result = await result
    if isinstance(result, str):
        return result
    if isinstance(result, (list, tuple)):
        return "\n\n".join(
            doc.text if hasattr(doc, "text") else str(doc)
            for doc in result
        )
    return str(result)


async def read_file_content(file_path: Path, get_parser: Optional[Callable] = None) -> str:
    """读取单个文件内容（支持 txt/md/pdf/docx）"""
    suffix = file_path.suffix.lower()

    if suffix in [".txt", ".md"]:
        return await asyncio.to_thread(
            lambda: file_path.read_text(encoding="utf-8")
        )

    if suffix == ".pdf":
        if get_parser is None:
            raise RuntimeError("解析 PDF 需要提供 get_parser 回调")
        parser = get_parser()
        return await parse_pdf_with_parser(file_path, parser)

    if suffix == ".docx":
        import docx
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])

    return ""


async def read_documents(data_path: str, get_parser: Optional[Callable] = None) -> List[str]:
    """读取目录下所有支持格式的文档"""
    dp = Path(data_path)
    if not dp.exists():
        raise FileNotFoundError(f"路径不存在: {dp}")

    documents = []
    for ext in ["*.txt", "*.md", "*.pdf", "*.docx"]:
        for fp in dp.rglob(ext):
            try:
                content = await read_file_content(fp, get_parser)
                if content:
                    documents.append(content)
                    print(f"[读取] {fp.name}")
            except Exception as e:
                print(f"[跳过] {fp.name}: {e}")
    return documents


class LightRAGSkill:

    def __init__(self, config: Optional[dict] = None, working_dir: Optional[str] = None):
        self._config = config or _load_skill_config()
        self._working_dir = working_dir or str(SKILL_DIR / "lightrag_db")
        self._rag: Optional[LightRAG] = None
        self._documents: List[str] = []

        # API 配置
        graphrag_cfg = self._config.get("graphrag", {})
        llm_cfg = graphrag_cfg.get("llm", {})
        embedding_cfg = graphrag_cfg.get("embedding", {})

        self._api_base = llm_cfg.get("api_base_url", "https://api-inference.modelscope.cn/v1")
        self._api_key = llm_cfg.get("api_key", "")
        self._llm_model = llm_cfg.get("api_model_id", "deepseek-ai/DeepSeek-V4-Flash")

        self._embed_provider = embedding_cfg.get("provider", "local")
        self._embed_model = embedding_cfg.get("model_name", "qwen3.5")
        self._embed_dim = embedding_cfg.get("embedding_dim", 1024)

        # 文档解析器配置
        parser_cfg = graphrag_cfg.get("document_parser", {})
        self._doc_parser_provider = parser_cfg.get("provider", "pypdf2")
        self._doc_parsers_cfg = parser_cfg.get("parsers", {})
        self._pdf_parser: Optional[Callable] = None

    def _load_pdf_parser(self) -> Callable:
        """从配置动态导入 PDF 解析器"""
        if self._pdf_parser is not None:
            return self._pdf_parser

        provider = self._doc_parser_provider
        cfg = self._doc_parsers_cfg.get(provider)
        if not cfg:
            raise RuntimeError(f"配置中未找到解析器 '{provider}'，可用: {list(self._doc_parsers_cfg.keys())}")

        # 确保 scripts 目录在 sys.path 中，模块才能被找到
        scripts_dir = str(SKILL_DIR / "scripts")
        if scripts_dir not in sys.path:
            sys.path.insert(0, scripts_dir)

        try:
            module = importlib.import_module(cfg["module"])
        except ModuleNotFoundError:
            # 如果直接导入失败，尝试从 scripts 目录导入
            module = importlib.import_module(cfg["module"].rsplit(".", 1)[-1])
        cls = getattr(module, cfg["class"])
        method = cfg.get("method", "")
        params = cfg.get("params", {})

        if method == "__call__" or not method:
            # 直接可调用函数/类
            parser = cls
        elif method:
            # 类 + 方法
            instance = cls(**params) if params else cls()
            parser = getattr(instance, method)
        else:
            # 类本身可调用
            parser = cls(**params) if params else cls()

        self._pdf_parser = parser
        return parser

    async def initialize(self):
        """异步初始化 LightRAG（带资源保护参数）"""
        print(f"[LightRAG] 正在初始化...")
        print(f"[LightRAG] LLM 并发限制: {LLM_SEMAPHORE._value}")
        print(f"[LightRAG] Embedding 并发限制: {EMBED_SEMAPHORE._value}")

        # 读取 lightrag 专用配置
        lightrag_cfg = self._config.get("graphrag", {}).get("lightrag", {})

        self._rag = LightRAG(
            working_dir=self._working_dir,
            # 资源保护参数
            max_parallel_insert=2,           # 文档级并发
            llm_model_max_async=2,           # LLM 并发，与信号量一致
            embedding_batch_num=8,           # 嵌入批处理大小
            default_llm_timeout=LLM_TIMEOUT,
            default_embedding_timeout=EMBED_TIMEOUT,
            addon_params={"insert_batch_size": 50},
            # LLM 和 Embedding 函数
            llm_model_func=self._safe_llm_call,
            embedding_func=EmbeddingFunc(
                embedding_dim=self._embed_dim,
                max_token_size=8192,
                func=self._safe_embedding_call
            ),
            # 从配置读取（不设则用 LightRAG 默认值）
            chunk_token_size=lightrag_cfg.get("chunk_token_size"),
            chunk_overlap_token_size=lightrag_cfg.get("chunk_overlap_token_size"),
        )

        # 必须调用：初始化存储后端
        await self._rag.initialize_storages()

        # 预加载本地 Embedding 模型（避免后续多线程并发加载导致 meta tensor 错误）
        if self._embed_provider == "local":
            from sentence_transformers import SentenceTransformer
            import torch
            torch.set_num_threads(4)
            embedding_cfg = self._config.get("graphrag", {}).get("embedding", {})
            local_path = embedding_cfg.get("local_model_path")
            model_name = embedding_cfg.get("model_name", "qwen3.5")
            model_path = local_path if local_path else model_name
            print(f"[LightRAG] 预加载 Embedding 模型: {model_path}")
            self._local_embed_model = SentenceTransformer(model_path, trust_remote_code=True)
            print(f"[LightRAG] Embedding 模型加载完成")

        print(f"[LightRAG] 初始化完成，工作目录: {self._working_dir}")

    async def finalize(self):
        """异步关闭存储连接"""
        if self._rag:
            await self._rag.finalize_storages()
            print("[LightRAG] 存储已安全关闭")

    async def _safe_llm_call(self, prompt: str, system_prompt: str | None = None,
                             history_messages: list[dict] | None = None, **kwargs) -> str:
        """带信号量的 LLM 调用。出错快速失败，提示调整并发数。"""
        async with LLM_SEMAPHORE:
            import openai
            client = openai.AsyncOpenAI(
                api_key=self._api_key,
                base_url=self._api_base,
                timeout=LLM_TIMEOUT,
                max_retries=0
            )

            # 构建 messages 数组
            messages = []
            if system_prompt:
                messages.append({"role": "system", "content": system_prompt})
            if history_messages:
                messages.extend(history_messages)
            messages.append({"role": "user", "content": prompt})

            # 过滤掉 LightRAG 内部参数，不传给 OpenAI API
            llm_kwargs = {k: v for k, v in kwargs.items()
                          if k not in ("hashing_kv", "_priority", "keyword_extraction", "enable_cot")}

            # DeepSeek 思考模式控制（从配置读取 enable_thinking）
            extra_body = {}
            enable_thinking = self._config.get("graphrag", {}).get("llm", {}).get("enable_thinking", True)
            if not enable_thinking:
                extra_body["thinking"] = {"type": "disabled"}

            try:
                create_kwargs = dict(
                    model=self._llm_model,
                    messages=messages,
                    **llm_kwargs
                )
                if extra_body:
                    create_kwargs["extra_body"] = extra_body
                response = await asyncio.wait_for(
                    client.chat.completions.create(**create_kwargs),
                    timeout=LLM_TIMEOUT
                )
                if not response.choices:
                    import json
                    print(f"[LLM] API 返回空 choices，完整响应: {json.dumps(response.model_dump(), indent=2, ensure_ascii=False)}")
                    raise RuntimeError(f"API 返回空 choices，疑似 ModelScope 兼容性问题")
                return response.choices[0].message.content
            except Exception as e:
                print(f"[LLM] 请求失败（当前并发 {LLM_SEMAPHORE._value}）: {e}")
                raise

    @async_retry(max_attempts=3, min_wait=1, max_wait=30)
    async def _safe_embedding_call(self, texts: List[str]) -> List[List[float]]:
        """带信号量的 Embedding 调用"""
        async with EMBED_SEMAPHORE:
            if self._embed_provider == "local":
                return await self._get_local_embeddings_async(texts)
            else:
                import openai
                client = openai.AsyncOpenAI(
                    api_key=self._api_key,
                    base_url=self._api_base,
                    timeout=EMBED_TIMEOUT,
                    max_retries=0
                )

                try:
                    response = await asyncio.wait_for(
                        client.embeddings.create(
                            model=self._embed_model,
                            input=texts
                        ),
                        timeout=EMBED_TIMEOUT
                    )
                    return [item.embedding for item in response.data]
                except asyncio.TimeoutError:
                    print(f"[Embedding] 请求超时 ({EMBED_TIMEOUT}s)")
                    raise
                except Exception as e:
                    print(f"[Embedding] 请求失败: {e}")
                    raise

    async def _get_local_embeddings_async(self, texts: List[str]) -> np.ndarray:
        """异步获取本地 Embedding（在线程池中执行）"""
        return await asyncio.to_thread(self._get_local_embeddings, texts)

    def _get_local_embeddings(self, texts: List[str]) -> np.ndarray:
        """本地 Embedding 模型 - 返回 numpy array（模型已在 initialize 时预加载）"""
        try:
            if not hasattr(self, '_local_embed_model') or self._local_embed_model is None:
                raise RuntimeError("Embedding 模型未初始化，请先调用 initialize()")
            embeddings = self._local_embed_model.encode(texts, convert_to_numpy=True)
            return embeddings
        except Exception as e:
            print(f"[LightRAG] 本地模型失败: {e}")
            import traceback
            traceback.print_exc()
            return np.array([])

    async def load_documents(self, data_path: Optional[str] = None) -> List[str]:
        """异步加载文档并构建索引"""
        start = time.time()
        if data_path is None:
            data_path = self._config.get("graphrag", {}).get("docs_dir", "./reference")

        documents = await read_documents(data_path, self._load_pdf_parser)
        if not documents:
            raise ValueError(f"未找到支持的文档文件: {data_path}")

        print(f"[LightRAG] 共加载 {len(documents)} 个文档，开始构建索引...")
        await self.ainsert(documents)

        print(f"[LightRAG] 索引构建完成，耗时 {time.time()-start:.2f}s")
        return documents

    async def ainsert(self, documents: List[str]):
        """增量插入文档内容到现有索引（直接封装 LightRAG.ainsert）"""
        if not self._rag:
            raise RuntimeError("LightRAG 未初始化")
        start = time.time()
        await self._rag.ainsert(documents)
        self._documents.extend(documents)
        print(f"[LightRAG] 增量插入 {len(documents)} 个文档，耗时 {time.time()-start:.2f}s")

    async def _read_file(self, file_path: Path) -> str:
        """异步读取文件"""
        return await read_file_content(file_path, self._load_pdf_parser)

    async def _parse_pdf(self, pdf_path: Path) -> str:
        """使用配置的解析器解析 PDF"""
        parser = self._load_pdf_parser()
        return await parse_pdf_with_parser(pdf_path, parser)

    def _read_docx(self, file_path: Path) -> str:
        """读取 Word 文档"""
        import docx
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])

    async def query(self, prompt: str, mode: str = "hybrid") -> Dict[str, Any]:
        """异步查询（带资源保护和引用）"""
        if not self._rag:
            raise RuntimeError("LightRAG 未初始化")

        print(f"[LightRAG] 查询 [{mode}]: {prompt[:50]}...")
        start = time.time()

        # 启用引用和上下文
        param = QueryParam(
            mode=mode,
            enable_rerank=False,  # 禁用 rerank 避免警告
            include_references=True,  # 启用引用
            only_need_context=False   # 返回完整答案
        )

        response = await self._rag.aquery(prompt, param=param)

        elapsed = time.time() - start
        print(f"[LightRAG] 查询完成，耗时 {elapsed:.2f}s")

        return {"answer": response, "mode": mode, "query": prompt, "elapsed_time": elapsed}


# =================================================================
# 便捷函数
# =================================================================

async def create_skill(config_path: Optional[str] = None) -> LightRAGSkill:
    """创建并初始化 Skill"""
    config = None
    if config_path:
        import yaml
        with open(config_path, "r", encoding="utf-8") as f:
            config = yaml.safe_load(f)

    skill = LightRAGSkill(config=config)
    await skill.initialize()
    return skill


async def load_and_index(data_path: str, config_path: Optional[str] = None) -> LightRAGSkill:
    """加载文档并构建索引"""
    skill = await create_skill(config_path)
    try:
        await skill.load_documents(data_path)
        return skill
    except Exception as e:
        await skill.finalize()
        raise e


async def query_index(skill: LightRAGSkill, prompt: str, mode: str = "hybrid") -> str:
    """查询索引"""
    result = await skill.query(prompt, mode=mode)
    return result["answer"]
